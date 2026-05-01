from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.services.docx_parser import parse_docx_to_blocks_and_structure


def test_parse_docx_to_blocks_and_structure_captures_styles(tmp_path: Path) -> None:
    doc_path = tmp_path / "styled.docx"
    doc = Document()

    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Styled text")
    run.bold = True
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x11, 0x22, 0x33)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Cell A"
    table.rows[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table.rows[0].cells[1].text = ""

    section = doc.sections[0]
    section.header.paragraphs[0].text = "Header X"
    section.footer.paragraphs[0].text = "Footer Y"
    doc.save(doc_path)

    blocks, structure = parse_docx_to_blocks_and_structure(doc_path, "DOCX1")
    assert blocks
    assert any(b.block_type == "table" for b in blocks)

    paragraph_blocks = [b for b in structure.blocks if b.block_type == "paragraph"]
    assert paragraph_blocks
    pb = paragraph_blocks[0]
    assert pb.paragraph_style is not None
    assert pb.paragraph_style.alignment is not None
    assert pb.run_styles
    assert pb.run_styles[0].bold is True
    assert pb.run_styles[0].font_name == "Calibri"
    assert pb.run_styles[0].font_size_pt is not None
    assert pb.run_styles[0].font_color_rgb == "112233"

    cell_blocks = [b for b in structure.blocks if b.block_type == "table_cell"]
    assert cell_blocks
    assert any(cb.cell_style is not None for cb in cell_blocks)
    assert any(cb.is_placeholder for cb in cell_blocks)

    assert any(b.block_type == "header_paragraph" and b.text == "Header X" for b in structure.blocks)
    assert any(b.block_type == "footer_paragraph" and b.text == "Footer Y" for b in structure.blocks)
