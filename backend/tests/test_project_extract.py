from pathlib import Path

import pytest
from docx import Document

from app.core.settings import Settings
from app.models.workspace import ResolvedProjectPaths
from app.services.extract_engine import (
    allocate_unique_doc_ids,
    list_input_docx_files,
    run_extract_overview,
    run_extract_overview_all_input_docs,
)


def test_list_input_docx_sorted(tmp_path: Path) -> None:
    wp = ResolvedProjectPaths(root=tmp_path)
    ind = wp.input_docs_dir
    ind.mkdir(parents=True)
    Document().save(ind / "B.docx")
    Document().save(ind / "a.docx")
    names = [p.name for p in list_input_docx_files(wp)]
    assert names == ["a.docx", "B.docx"]


def test_allocate_unique_doc_ids_collision(tmp_path: Path) -> None:
    paths = [
        tmp_path / "same.docx",
        tmp_path / "sub" / "same.docx",
    ]
    paths[1].parent.mkdir(parents=True)
    ids = allocate_unique_doc_ids(paths)
    assert ids == ["same", "same_1"]


def test_run_extract_overview_all_batches_input_docs(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    wd = ResolvedProjectPaths(root=tmp_path)
    ind = wd.input_docs_dir
    ind.mkdir(parents=True)
    Document().save(ind / "second.docx")
    Document().save(ind / "first.docx")

    settings = Settings(
        openai_api_key="sk-test",
        project_paths=wd,
    )
    touched: list[str] = []

    def fake(**kwargs):  # type: ignore[no-untyped-def]
        touched.append(f"{kwargs['input_doc_path'].name}:{kwargs['doc_id']}")
        return {"ok": 1, "doc_id": kwargs["doc_id"]}

    monkeypatch.setattr("app.services.extract_engine.run_extract_overview", fake)

    out = run_extract_overview_all_input_docs(settings=settings)
    assert out["count"] == 2
    assert touched == ["first.docx:first", "second.docx:second"]


def test_run_extract_overview_all_errors_when_input_docs_empty(tmp_path: Path) -> None:
    wd = ResolvedProjectPaths(root=tmp_path)
    wd.input_docs_dir.mkdir(parents=True)
    settings = Settings(openai_api_key="sk", project_paths=wd)
    with pytest.raises(ValueError, match=r"No \.docx files"):
        run_extract_overview_all_input_docs(settings=settings)


def test_run_extract_overview_requires_api_key(tmp_path: Path) -> None:
    doc_path = tmp_path / "a.docx"
    Document().save(doc_path)
    settings = Settings(
        openai_api_key="",
        project_paths=ResolvedProjectPaths(root=tmp_path),
    )
    with pytest.raises(ValueError, match="OPENAI_API_KEY"):
        run_extract_overview(
            input_doc_path=doc_path,
            doc_id="D1",
            settings=settings,
        )


def test_run_extract_overview_missing_input(tmp_path: Path) -> None:
    settings = Settings(
        openai_api_key="sk-test",
        project_paths=ResolvedProjectPaths(root=tmp_path),
    )
    missing = tmp_path / "nope.docx"
    with pytest.raises(FileNotFoundError, match="Input document not found"):
        run_extract_overview(
            input_doc_path=missing,
            doc_id="D1",
            settings=settings,
        )
