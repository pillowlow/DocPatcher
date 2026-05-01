from pathlib import Path

from docx import Document

from app.services.docx_parser import parse_docx_to_blocks


def test_parse_docx_to_blocks_smoke(tmp_path: Path) -> None:
    doc_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("")
    doc.add_paragraph("Second paragraph.")
    doc.save(doc_path)

    blocks = parse_docx_to_blocks(doc_path, doc_id="DOC001")
    assert len(blocks) == 2
    assert blocks[0].block_id == "DOC001-B0000"
    assert blocks[1].text == "Second paragraph."
