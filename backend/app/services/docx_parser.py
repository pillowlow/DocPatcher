from pathlib import Path
from typing import Callable
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell
from docx.table import Table
from docx.text.run import Run
from docx.text.paragraph import Paragraph

from app.models.block import Block, BlockPosition
from app.models.doc_structure import (
    DocCellStyle,
    DocParagraphStyle,
    DocRunStyle,
    DocStructureBlock,
    DocStructureDocument,
    DocStructurePosition,
    DocStructureXmlStats,
)


def _to_pt(value: object) -> float | None:
    if value is None:
        return None
    pt = getattr(value, "pt", None)
    if pt is None:
        return None
    return float(pt)


def _underline_to_bool(value: object) -> bool | None:
    if value is None:
        return None
    if isinstance(value, bool):
        return value
    return True


def _run_style(run: Run) -> DocRunStyle:
    color_rgb = None
    if run.font.color is not None and run.font.color.rgb is not None:
        color_rgb = str(run.font.color.rgb)
    highlight = None
    if run.font.highlight_color is not None:
        highlight = str(run.font.highlight_color)
    return DocRunStyle(
        text=run.text,
        font_name=run.font.name,
        font_size_pt=_to_pt(run.font.size),
        bold=run.bold,
        italic=run.italic,
        underline=_underline_to_bool(run.underline),
        font_color_rgb=color_rgb,
        highlight_color=highlight,
    )


def _paragraph_style(para: Paragraph) -> DocParagraphStyle:
    fmt = para.paragraph_format
    line_spacing = fmt.line_spacing
    numeric_line_spacing = None
    if isinstance(line_spacing, (int, float)):
        numeric_line_spacing = float(line_spacing)
    return DocParagraphStyle(
        style_name=para.style.name if para.style is not None else None,
        alignment=str(para.alignment) if para.alignment is not None else None,
        space_before_pt=_to_pt(fmt.space_before),
        space_after_pt=_to_pt(fmt.space_after),
        line_spacing=numeric_line_spacing,
        left_indent_pt=_to_pt(fmt.left_indent),
        right_indent_pt=_to_pt(fmt.right_indent),
        first_line_indent_pt=_to_pt(fmt.first_line_indent),
    )


def _cell_style(cell: _Cell, table: Table) -> DocCellStyle:
    shading_fill = None
    tc_pr = cell._tc.tcPr
    if tc_pr is not None:
        shd = tc_pr.find(qn("w:shd"))
        if shd is not None:
            shading_fill = shd.get(qn("w:fill"))
    return DocCellStyle(
        vertical_alignment=(
            str(cell.vertical_alignment) if cell.vertical_alignment is not None else None
        ),
        shading_fill=shading_fill,
        table_style=(table.style.name if table.style is not None else None),
    )


def analyze_docx_xml_parts(docx_path: Path) -> DocStructureXmlStats:
    xml_infos: list[tuple[int, int]] = []
    with ZipFile(docx_path) as zf:
        for info in zf.infolist():
            if not info.filename.lower().endswith(".xml"):
                continue
            if not info.filename.lower().startswith("word/"):
                continue
            xml_infos.append((info.compress_size, info.file_size))
    return DocStructureXmlStats(
        xml_parts_total=len(xml_infos),
        xml_parts_compressed_bytes=sum(i[0] for i in xml_infos),
        xml_parts_uncompressed_bytes=sum(i[1] for i in xml_infos),
    )


def parse_docx_to_blocks_and_structure(
    docx_path: Path,
    doc_id: str,
    *,
    on_xml_part_done: Callable[[str], None] | None = None,
) -> tuple[list[Block], DocStructureDocument]:
    document = Document(docx_path)
    xml_stats = analyze_docx_xml_parts(docx_path)
    blocks: list[Block] = []
    structure_blocks: list[DocStructureBlock] = []

    order = 0
    para_idx = 0

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            para = Paragraph(child, document)
            text = para.text.strip()
            placeholder = not bool(text)
            run_styles = [_run_style(run) for run in para.runs]
            structure_blocks.append(
                DocStructureBlock(
                    structure_id=f"{doc_id}-S{order:05d}",
                    source_part="word/document.xml",
                    block_type="blank_paragraph" if placeholder else "paragraph",
                    order_index=order,
                    text=text if text else "<BLANK_PARAGRAPH>",
                    is_placeholder=placeholder,
                    position=DocStructurePosition(paragraph_index=para_idx),
                    paragraph_style=_paragraph_style(para),
                    run_styles=run_styles,
                )
            )
            if text:
                blocks.append(
                    Block(
                        block_id=f"{doc_id}-B{para_idx:04d}",
                        doc_id=doc_id,
                        file_name=docx_path.name,
                        text=text,
                        position=BlockPosition(paragraph_index=para_idx),
                    )
                )
            para_idx += 1
            order += 1
            continue

        if isinstance(child, CT_Tbl):
            table = Table(child, document)
            row_texts: list[str] = []
            for r_idx, row in enumerate(table.rows):
                cell_texts: list[str] = []
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if not cell_text:
                        cell_text = "<BLANK_CELL>"
                    cell_run_styles: list[DocRunStyle] = []
                    first_para_style = None
                    for para in cell.paragraphs:
                        if first_para_style is None:
                            first_para_style = _paragraph_style(para)
                        cell_run_styles.extend(_run_style(run) for run in para.runs)
                    cell_texts.append(cell_text)
                    structure_blocks.append(
                        DocStructureBlock(
                            structure_id=f"{doc_id}-S{order:05d}",
                            source_part="word/document.xml",
                            block_type="table_cell",
                            order_index=order,
                            text=cell_text,
                            is_placeholder=(cell_text == "<BLANK_CELL>"),
                            position=DocStructurePosition(
                                row_index=r_idx,
                                col_index=c_idx,
                            ),
                            paragraph_style=first_para_style,
                            cell_style=_cell_style(cell, table),
                            run_styles=cell_run_styles,
                        )
                    )
                    order += 1
                row_texts.append(" | ".join(cell_texts))
            table_text = "\n".join(row_texts).strip()
            if table_text:
                blocks.append(
                    Block(
                        block_id=f"{doc_id}-B{para_idx:04d}",
                        doc_id=doc_id,
                        file_name=docx_path.name,
                        block_type="table",
                        text=table_text,
                        position=BlockPosition(paragraph_index=para_idx),
                    )
                )
                para_idx += 1
            continue

    seen_partnames: set[str] = {"word/document.xml"}
    if on_xml_part_done:
        on_xml_part_done("word/document.xml")

    for s_idx, section in enumerate(document.sections):
        for part_obj, part_kind in ((section.header.part, "header"), (section.footer.part, "footer")):
            partname = str(part_obj.partname).lstrip("/")
            if partname in seen_partnames:
                continue
            seen_partnames.add(partname)
            paragraphs = part_obj.element.xpath(".//w:p")
            for p_idx, p_elem in enumerate(paragraphs):
                para = Paragraph(p_elem, document)
                text = para.text.strip()
                placeholder = not bool(text)
                run_styles = [_run_style(run) for run in para.runs]
                structure_blocks.append(
                    DocStructureBlock(
                        structure_id=f"{doc_id}-S{order:05d}",
                        source_part=partname,
                        block_type=f"{part_kind}_paragraph",
                        order_index=order,
                        text=text if text else f"<BLANK_{part_kind.upper()}_PARAGRAPH>",
                        is_placeholder=placeholder,
                        position=DocStructurePosition(
                            section_index=s_idx,
                            paragraph_index=p_idx,
                        ),
                        paragraph_style=_paragraph_style(para),
                        run_styles=run_styles,
                    )
                )
                order += 1
                if text:
                    blocks.append(
                        Block(
                            block_id=f"{doc_id}-B{para_idx:04d}",
                            doc_id=doc_id,
                            file_name=docx_path.name,
                            block_type=f"{part_kind}_paragraph",
                            text=text,
                            position=BlockPosition(paragraph_index=para_idx),
                        )
                    )
                    para_idx += 1
            if on_xml_part_done:
                on_xml_part_done(partname)

    structure_doc = DocStructureDocument(
        doc_id=doc_id,
        file_name=docx_path.name,
        xml_stats=xml_stats,
        blocks=structure_blocks,
    )
    return blocks, structure_doc


def parse_docx_to_blocks(docx_path: Path, doc_id: str) -> list[Block]:
    blocks, _ = parse_docx_to_blocks_and_structure(docx_path, doc_id)
    return blocks
